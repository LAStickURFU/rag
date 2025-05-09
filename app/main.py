# Копируем содержимое app/main.py с изменением метода upload_document
import os
import logging
import asyncio
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any, Union, Tuple
from dotenv import load_dotenv
import re
import secrets
import uuid
import mimetypes
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from striprtf.striprtf import rtf_to_text
from bs4 import BeautifulSoup
import yaml
import json
from odf.opendocument import load as odf_load
from odf.text import P
from ebooklib import epub
import time
import math

from fastapi import FastAPI, Depends, HTTPException, status, Request, Form, UploadFile, File, BackgroundTasks
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from jose import JWTError, jwt
from passlib.context import CryptContext
import torch
from sqlalchemy.orm import Session
from pydantic import BaseModel
import glob
from pathlib import Path
import requests
from datasets import load_dataset

from app.logging_config import setup_logging
from app.database import get_db, create_tables, engine
from app.models import User as UserModel, Chat as ChatModel, Document as DocumentModel, ModelConfig as ModelConfigModel
from app.schemas import UserCreate, User, Token, TokenData, Chat, ChatMessage, QuestionRequest, ModelSettings, ChatResponse, ChunkInfo, RagasEvaluationRequest, RagasEvalItem
from app.ollama_client import get_ollama_instance, preload_model
from app.rag import RAGService, Document
from app.config import EMBEDDING_MODEL, CHUNK_SIZE, CHUNK_OVERLAP, USE_HYBRID_SEARCH, USE_RERANKER, DENSE_WEIGHT, RERANKER_WEIGHT, USE_ADAPTIVE_K, CROSS_ENCODER_MODEL, LANGUAGE, SPACY_MODEL
from scripts.evaluate_rag import RAGEvaluator, get_builtin_datasets

# Настройка логирования через конфигурацию
setup_logging()
logger = logging.getLogger(__name__)

# Загрузка переменных окружения
load_dotenv()

# Конфигурация безопасности
SECRET_KEY = os.getenv("JWT_SECRET_KEY", "your-secret-key")
ALGORITHM = os.getenv("JWT_ALGORITHM", "HS256")
ACCESS_TOKEN_EXPIRE_MINUTES = 30

# Инициализация хеширования паролей
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# Инициализация FastAPI приложения
app = FastAPI(title="RAG Service")

# Инициализация таблиц в базе данных при запуске приложения
create_tables()

# Добавление CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Для продакшена нужно указать конкретные домены
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Healthcheck эндпоинт
@app.get("/api/healthcheck")
async def healthcheck():
    """
    Проверка здоровья сервиса, возвращает информацию о статусе различных компонентов
    """
    try:
        # Проверяем доступность Ollama
        ollama_available = False
        ollama_model = os.getenv("OLLAMA_MODEL", "mistral:7b-instruct")
        try:
            ollama = get_ollama_instance()
            ollama_available = await ollama.check_model_availability()
        except Exception as e:
            logger.warning(f"Ollama недоступен: {str(e)}")
        
        # Проверяем состояние службы RAG
        rag_service_available = False
        try:
            # Пробуем получить службу RAG
            rag = get_rag_service()
            rag_service_available = True
        except Exception as e:
            logger.warning(f"Служба RAG недоступна: {str(e)}")
        
        # Собираем информацию о среде
        environment_info = {
            "embedding_model": os.getenv("EMBEDDING_MODEL", "intfloat/multilingual-e5-base"),
            "spacy_model": os.getenv("SPACY_MODEL", "ru_core_news_md"),
            "chunk_size": int(os.getenv("CHUNK_SIZE", "400")),
            "chunk_overlap": int(os.getenv("CHUNK_OVERLAP", "100")),
            "python_version": os.getenv("PYTHON_VERSION", ""),
            "pytorch_version": torch.__version__ if 'torch' in globals() else "not available"
        }
        
        # Возвращаем полный отчет о состоянии
        return {
            "status": "ok",
            "timestamp": datetime.utcnow().isoformat(),
            "components": {
                "rag_service": "available" if rag_service_available else "unavailable",
                "ollama": "available" if ollama_available else "unavailable",
                "ollama_model": ollama_model
            },
            "environment": environment_info
        }
    except Exception as e:
        logger.error(f"Ошибка при проверке состояния: {str(e)}")
        return {
            "status": "error",
            "timestamp": datetime.utcnow().isoformat(),
            "error": str(e)
        }

# Глобальные переменные
rag_service = None
model_initialized = False

async def initialize_ollama():
    """Инициализирует клиент Ollama и пытается загрузить модель."""
    try:
        model_name = os.getenv("OLLAMA_MODEL", "mistral:7b-instruct")
        logger.info(f"Initializing Ollama client with model {model_name}")
        
        # Проверяем наличие Ollama
        ollama_client = get_ollama_instance(model_name)
        available = await ollama_client.check_model_availability()
        
        if not available:
            # Пытаемся автоматически запустить Ollama если она установлена
            import platform
            import subprocess
            
            system = platform.system()
            logger.info(f"Attempting to automatically start Ollama on {system}")
            
            try:
                if system == "Darwin":  # macOS
                    # Пробуем запустить через open
                    subprocess.Popen(["open", "-a", "Ollama"])
                    logger.info("Attempted to start Ollama app on macOS")
                elif system == "Windows":
                    # Пытаемся запустить исполняемый файл Ollama
                    subprocess.Popen(["ollama", "serve"])
                elif system == "Linux":
                    # Пытаемся запустить ollama serve
                    subprocess.Popen(["ollama", "serve"])
                
                # Ждем немного, чтобы Ollama успела запуститься
                import asyncio
                logger.info("Waiting for Ollama to start...")
                await asyncio.sleep(5)
                
                # Проверяем снова доступность
                available = await ollama_client.check_model_availability()
                if available:
                    logger.info("Successfully started Ollama")
                else:
                    logger.warning("Could not automatically start Ollama")
            except Exception as e:
                logger.error(f"Error starting Ollama: {str(e)}")
        
        await ollama_client.ensure_model_loaded()
        return ollama_client
    except Exception as e:
        logger.error(f"Error initializing Ollama: {str(e)}")
        return None

def get_rag_service():
    global rag_service
    if rag_service is None:
        try:
            logger.info("Initializing real RAGService")
            rag_service = RAGService(
                index_name="default",
                model_name=EMBEDDING_MODEL,
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
                use_hybrid_search=USE_HYBRID_SEARCH,
                use_reranker=USE_RERANKER,
                dense_weight=DENSE_WEIGHT,
                reranker_weight=RERANKER_WEIGHT,
                use_adaptive_k=USE_ADAPTIVE_K,
                cross_encoder_model=CROSS_ENCODER_MODEL,
                language=LANGUAGE,
                spacy_model=SPACY_MODEL
            )
            logger.info("RAGService initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing RAGService: {str(e)}")
            raise
    return rag_service

# Функции безопасности
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def get_user(db: Session, username: str):
    return db.query(UserModel).filter(UserModel.username == username).first()

def authenticate_user(db: Session, username: str, password: str):
    user = get_user(db, username)
    if not user:
        return False
    if not verify_password(password, user.hashed_password):
        return False
    return user

def create_access_token(data: dict, expires_delta: timedelta = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
        token_data = TokenData(username=username)
    except JWTError:
        raise credentials_exception
    user = get_user(db, username=token_data.username)
    if user is None:
        raise credentials_exception
    return user

# Модели запросов
class DocumentUploadRequest(BaseModel):
    title: str
    source: str = "manual"

# Событие при запуске приложения
@app.on_event("startup")
async def startup_event():
    """Выполняется при запуске приложения."""
    # Запускаем инициализацию модели в отдельном потоке, чтобы не блокировать запуск сервера
    from concurrent.futures import ThreadPoolExecutor
    
    executor = ThreadPoolExecutor()
    logger.info("Starting application...")
    
    # Инициализируем клиент Ollama асинхронно
    try:
        await initialize_ollama()
        logger.info("Ollama initialization process started")
    except Exception as e:
        logger.error(f"Error starting Ollama initialization: {str(e)}")
    # Автоматическая загрузка датасетов для оценки
    try:
        get_builtin_datasets()
    except Exception as e:
        logger.error(f"Ошибка автозагрузки датасетов: {e}")

# Эндпоинты
@app.post("/token", response_model=Token)
async def login_for_access_token(
    form_data: OAuth2PasswordRequestForm = Depends(),
    db: Session = Depends(get_db)
):
    user = authenticate_user(db, form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Неверное имя пользователя или пароль",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/register", response_model=User)
async def register_user(user: UserCreate, db: Session = Depends(get_db)):
    db_user = get_user(db, username=user.username)
    if db_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Пользователь уже существует"
        )
    hashed_password = get_password_hash(user.password)
    db_user = UserModel(username=user.username, hashed_password=hashed_password)
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    return db_user

def get_user_model_config(user_id: int, db: Session) -> ModelConfigModel:
    """Возвращает конфигурацию модели для пользователя или создает её по умолчанию"""
    config = db.query(ModelConfigModel).filter(ModelConfigModel.user_id == user_id).first()
    
    if not config:
        # Создаем конфигурацию по умолчанию
        config = ModelConfigModel(user_id=user_id)
        db.add(config)
        db.commit()
        db.refresh(config)
    
    return config

@app.post("/ask", response_model=ChatResponse)
async def ask_question(
    question_request: QuestionRequest,
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Обработка вопроса с использованием RAG.
    """
    try:
        logger.info(f"Processing question from user {current_user.username}: {question_request.question}")
        logger.debug("Обработка вопроса от пользователя %s: %s", current_user.username, question_request.question)
        
        # Получаем настройки пользователя
        logger.debug("Получение настроек пользователя")
        user_config = get_user_model_config(current_user.id, db)
        
        # Получаем клиент Ollama
        logger.debug("Получение Ollama клиента")
        ollama_client = get_ollama_instance(user_config.model_name)
        
        # Получаем сервис RAG
        logger.debug("Получение RAG-сервиса")
        rag_service = get_rag_service()
        
        # Метрики времени
        t0 = time.perf_counter()
        # Поиск релевантных чанков
        logger.debug("Поиск релевантных чанков")
        relevant_chunks = rag_service.search(question_request.question, top_k=5)
        t1 = time.perf_counter()
        retrieval_time_ms = int((t1 - t0) * 1000)
        
        # Формируем запрос с контекстом
        logger.debug("Найдено %d релевантных чанков", len(relevant_chunks))
        logger.debug("Генерация промпта с контекстом")
        prompt = rag_service.generate_prompt(question_request.question, top_k_chunks=5)
        logger.info("Generated RAG prompt with context")
        logger.debug("Промпт с контекстом сгенерирован")
        
        # Запрос к модели
        logger.debug("Запрос к Ollama на генерацию ответа")
        t2 = time.perf_counter()
        try:
            response = await ollama_client.generate(prompt, temperature=user_config.temperature)
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            # Если произошла ошибка, возвращаем сообщение об ошибке
            if "Failed to connect to Ollama" in str(e):
                response = f"Ошибка: не удалось подключиться к Ollama. Пожалуйста, установите и запустите Ollama, следуя инструкции на https://ollama.com/download"
            else:
                response = f"Произошла ошибка при генерации ответа: {str(e)}"
        t3 = time.perf_counter()
        generation_time_ms = int((t3 - t2) * 1000)
        
        logger.info("Received response from Ollama")
        logger.debug("Получен ответ от Ollama")
        
        # Сохраняем результат в БД
        logger.info("Saving chat to database")
        logger.debug("Сохранение чата в базу данных")
        
        # Преобразуем результаты поиска в формат для сохранения
        relevant_chunks_data = []
        for chunk, score in relevant_chunks:
            # Преобразуем метаданные в словарь для сериализации
            metadata = {k: v for k, v in chunk.metadata.items() if k != 'embedding'}
            
            relevant_chunks_data.append({
                "text": chunk.text,
                "relevance": float(score),
                "doc_id": chunk.doc_id,
                "chunk_id": chunk.chunk_id,
                "metadata": metadata
            })
        
        # Собираем метаинформацию
        meta = {
            "retrieval_time_ms": retrieval_time_ms,
            "generation_time_ms": generation_time_ms,
            "retrieved_chunks_count": len(relevant_chunks),
            "prompt_length": len(prompt) if prompt else None,
            "response_length": len(response) if response else None,
            "model": getattr(user_config, 'model_name', None),
            "temperature": getattr(user_config, 'temperature', None),
            "top_k_chunks": 5,
            "trace_id": str(uuid.uuid4()),
            "timestamp": datetime.utcnow().isoformat(),
            "rag_used": True
        }
        
        # Создаем новую запись чата
        chat = ChatModel(
            user_id=current_user.id,
            question=question_request.question,
            response=response,
            relevant_chunks=relevant_chunks_data,
            rag_used=meta.get('rag_used', True)
        )
        
        db.add(chat)
        db.commit()
        db.refresh(chat)
        
        logger.info("Chat saved successfully")
        logger.debug("Чат успешно сохранен")
        
        # Преобразуем модель БД в модель ответа
        return ChatResponse(
            id=chat.id,
            user_id=chat.user_id,
            question=chat.question,
            response=chat.response,
            created_at=chat.created_at,
            relevant_chunks=chat.relevant_chunks,
            meta=meta,
            rag_used=chat.rag_used
        )
    except Exception as e:
        logger.error(f"Error processing question: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка обработки запроса: {str(e)}"
        )

@app.get("/chats", response_model=List[ChatResponse])
async def get_chat_history(
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    chats = db.query(ChatModel).filter(
        ChatModel.user_id == current_user.id
    ).order_by(ChatModel.created_at.desc()).all()
    
    return chats

@app.delete("/chats/clear")
async def clear_chat_history(
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Очищает всю историю чата пользователя.
    """
    db.query(ChatModel).filter(ChatModel.user_id == current_user.id).delete()
    db.commit()
    return {"status": "success", "message": "История чата очищена"}

# Универсальный парсер текстовых форматов

def extract_text(file: UploadFile, contents: bytes) -> str:
    ext = file.filename.lower().split('.')[-1]
    if ext in ['txt', 'md', 'log', 'tex', 'jsonl']:
        return contents.decode('utf-8', errors='ignore')
    elif ext == 'pdf':
        from io import BytesIO
        reader = PdfReader(BytesIO(contents))
        return '\n'.join(page.extract_text() or '' for page in reader.pages)
    elif ext in ['doc', 'docx']:
        from io import BytesIO
        doc = DocxDocument(BytesIO(contents))
        return '\n'.join([p.text for p in doc.paragraphs])
    elif ext == 'rtf':
        return rtf_to_text(contents.decode('utf-8', errors='ignore'))
    elif ext in ['json']:
        try:
            obj = json.loads(contents.decode('utf-8', errors='ignore'))
            return json.dumps(obj, ensure_ascii=False, indent=2)
        except Exception:
            return contents.decode('utf-8', errors='ignore')
    elif ext in ['yaml', 'yml']:
        try:
            obj = yaml.safe_load(contents.decode('utf-8', errors='ignore'))
            return yaml.dump(obj, allow_unicode=True)
        except Exception:
            return contents.decode('utf-8', errors='ignore')
    elif ext in ['csv', 'tsv']:
        return contents.decode('utf-8', errors='ignore')
    elif ext in ['html', 'htm']:
        soup = BeautifulSoup(contents, 'html.parser')
        return soup.get_text()
    elif ext == 'odt':
        from io import BytesIO
        doc = odf_load(BytesIO(contents))
        texts = []
        for elem in doc.getElementsByType(P):
            texts.append(str(elem))
        return '\n'.join(texts)
    elif ext == 'epub':
        from io import BytesIO
        book = epub.read_epub(BytesIO(contents))
        texts = []
        for item in book.get_items():
            if item.get_type() == epub.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                texts.append(soup.get_text())
        return '\n'.join(texts)
    elif ext == 'xml':
        return contents.decode('utf-8', errors='ignore')
    elif ext == 'ini':
        return contents.decode('utf-8', errors='ignore')
    else:
        raise HTTPException(status_code=400, detail=f'Формат {ext} не поддерживается')

# Изменённый эндпоинт загрузки документов
@app.post("/documents/upload")
async def upload_documents(
    files: List[UploadFile] = File(...),
    titles: List[str] = Form(...),
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    results = []
    for idx, file in enumerate(files):
        title = titles[idx] if idx < len(titles) else file.filename.split(".")[0]
        source = "manual"  # всегда ручной ввод для этого endpoint
        file_name = file.filename
        file_size = file.size if hasattr(file, 'size') else None
        doc_uuid = str(uuid.uuid4())
        now = datetime.utcnow()
        db_document = DocumentModel(
            title=title,
            content="",  # временно пусто, заполним после извлечения
            source=source,
            user_id=current_user.id,
            status="uploaded",
            file_name=file_name,
            file_size=file_size,
            uuid=doc_uuid,
            created_at=now,
            updated_at=now,
            chunks_count=0,
            error_message=None
        )
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        try:
            db_document.status = "indexing"
            db.commit()
            contents = await file.read()
            text = extract_text(file, contents)
            db_document.content = text
            rag = get_rag_service()
            document = Document(
                content=text,
                metadata={"title": title, "source": source, "user_id": current_user.id}
            )
            chunk_ids = rag.add_document(document, doc_id=str(db_document.id))
            db_document.chunks_count = len(chunk_ids)
            db_document.status = "indexed"
            db_document.updated_at = datetime.utcnow()
            db_document.error_message = None
            db.commit()
            results.append({
                "id": db_document.id,
                "title": db_document.title,
                "chunks_indexed": len(chunk_ids),
                "created_at": db_document.created_at,
                "status": db_document.status
            })
        except Exception as e:
            db_document.status = "error"
            db_document.error_message = str(e)
            db_document.updated_at = datetime.utcnow()
            db.commit()
            results.append({"error": str(e), "filename": file.filename, "id": db_document.id})
    return {"results": results}

@app.get("/documents")
async def get_documents(
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    documents = db.query(DocumentModel).filter(
        DocumentModel.user_id == current_user.id
    ).order_by(DocumentModel.created_at.desc()).all()
    
    return [
        {
            "id": doc.id,
            "uuid": doc.uuid,
            "title": doc.title,
            "source": doc.source,
            "created_at": doc.created_at,
            "updated_at": doc.updated_at,
            "status": doc.status,
            "file_name": doc.file_name,
            "file_size": doc.file_size,
            "chunks_count": doc.chunks_count,
            "error_message": doc.error_message
        } for doc in documents
    ]

@app.post("/documents/reindex")
async def reindex_documents(
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        documents = db.query(DocumentModel).filter(
            DocumentModel.user_id == current_user.id
        ).order_by(DocumentModel.created_at).all()
        if not documents:
            return {"status": "success", "message": "Документы для индексации не найдены", "indexed": 0}
        rag = get_rag_service()
        # Фикс для MockRAGService: если нет rag.index.clear(), очищаем rag.chunks
        if hasattr(rag, 'index') and hasattr(rag.index, 'clear'):
            rag.index.clear()
        elif hasattr(rag, 'chunks'):
            rag.chunks.clear()
        total_chunks = 0
        indexed_docs = 0
        for doc in documents:
            document = Document(
                content=doc.content, 
                metadata={"title": doc.title, "source": doc.source, "user_id": current_user.id}
            )
            chunk_ids = rag.add_document(document, doc_id=str(doc.id))
            total_chunks += len(chunk_ids)
            indexed_docs += 1
            await asyncio.sleep(0.1)
        return {
            "status": "success", 
            "message": "Документы успешно переиндексированы",
            "indexed_documents": indexed_docs,
            "total_chunks": total_chunks
        }
    except Exception as e:
        logger.error(f"Error reindexing documents: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Ошибка переиндексации документов: {str(e)}")

@app.get("/model/settings", response_model=ModelSettings)
async def get_model_settings(
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Получение текущих настроек модели для пользователя"""
    config = get_user_model_config(current_user.id, db)
    
    return ModelSettings(
        temperature=config.temperature,
        top_p=config.top_p,
        max_tokens=config.max_tokens,
        top_k_chunks=config.top_k_chunks,
        context_window=config.context_window
    )

@app.post("/model/settings", response_model=ModelSettings)
async def update_model_settings(
    settings: ModelSettings,
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Обновление настроек модели для пользователя"""
    config = get_user_model_config(current_user.id, db)
    
    # Валидация и ограничение значений
    config.temperature = max(0.0, min(1.0, settings.temperature))
    config.top_p = max(0.1, min(1.0, settings.top_p))
    config.max_tokens = max(100, min(8192, settings.max_tokens))
    config.top_k_chunks = max(1, min(20, settings.top_k_chunks))
    config.context_window = max(1024, min(16384, settings.context_window))
    
    db.commit()
    db.refresh(config)
    
    return ModelSettings(
        temperature=config.temperature,
        top_p=config.top_p,
        max_tokens=config.max_tokens,
        top_k_chunks=config.top_k_chunks,
        context_window=config.context_window
    )

@app.post("/direct-ask")
async def direct_ask(
    question_request: QuestionRequest,
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Напрямую задает вопрос LLM модели без использования RAG.
    Используется в качестве запасного варианта, если основной RAG не работает.
    """
    question = question_request.question
    logger.debug("Прямой запрос от пользователя %s: %s", current_user.username, question)
    
    try:
        # Получаем настройки пользователя
        user_config = get_user_model_config(current_user.id, db)
        
        # Получаем Ollama клиент
        ollama = get_ollama_instance(os.getenv("OLLAMA_MODEL", "mistral"))
        
        # Формируем простой промпт
        prompt = f"""Ответь на следующий вопрос пользователя:

Вопрос: {question}

Ответ:"""
        
        # Генерируем ответ
        generation_options = {
            'temperature': user_config.temperature,
            'num_predict': user_config.max_tokens,
            'top_p': user_config.top_p,
            'stop': ['<|im_end|>', '</answer>', '\n\n\n'] 
        }
        
        response = await ollama.generate(prompt, **generation_options)
        
        # Очищаем ответ от возможных инструкций
        if "Ответ:" in response:
            response = response.split("Ответ:")[-1].strip()
        
        # Сохраняем в истории
        chat = ChatModel(
            user_id=current_user.id,
            question=question,
            response=response,
            relevant_chunks=[],  # Нет релевантных чанков, т.к. не используем RAG
            rag_used=False
        )
        db.add(chat)
        db.commit()
        db.refresh(chat)
        
        # Создаем ответ
        chat_response = ChatResponse(
            id=chat.id,
            user_id=chat.user_id,
            question=chat.question,
            response=chat.response,
            created_at=chat.created_at,
            relevant_chunks=[],  # Пустой список, т.к. не используем RAG
            rag_used=False
        )
        
        return chat_response
    except Exception as e:
        logger.debug("Ошибка прямого запроса: %s", str(e))
        logger.debug("Трассировка: %s", traceback.format_exc())
        logger.error(f"Error in direct ask: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Ошибка обработки запроса: {str(e)}")

# Модели данных для оценки RAG
class RAGEvaluationItem(BaseModel):
    question: str
    answer: str
    ground_truth: str


class RAGEvaluationRequest(BaseModel):
    eval_items: List[RAGEvaluationItem] = []
    description: Optional[str] = None


class RAGEvaluationResponse(BaseModel):
    metrics: Dict[str, float]
    summary: str
    evaluation_file: str

# Глобальные словари для прогресса и логов
rag_eval_progress: Dict[str, Dict] = {}
rag_eval_logs: Dict[str, list] = {}

# Фоновая задача для оценки RAG

def run_rag_evaluation_task(task_id, request, db, user_id):
    import traceback
    import json
    import platform
    import importlib
    import sys
    from datetime import datetime
    
    start_time = datetime.now()
    
    try:
        # Инициализация прогресса и логов
        rag_eval_progress[task_id] = {
            "status": "initializing", 
            "progress": 0,
            "start_time": start_time.isoformat(),
            "user_id": user_id
        }
        rag_eval_logs[task_id] = []
        
        log_entry = f"[{datetime.now().isoformat()}] Старт оценки RAG (ID: {task_id})..."
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Логируем информацию о запросе для отладки
        try:
            request_info = {
                "eval_items_count": len(request.eval_items),
                "has_items": bool(request.eval_items),
                "description": request.description,
                "sample_item": request.eval_items[0].dict() if request.eval_items else None
            }
            log_entry = f"[{datetime.now().isoformat()}] Информация о запросе: {json.dumps(request_info, ensure_ascii=False)}"
            rag_eval_logs[task_id].append(log_entry)
            logger.info(log_entry)
        except Exception as e:
            log_entry = f"[{datetime.now().isoformat()}] Ошибка при логировании запроса: {e}"
            rag_eval_logs[task_id].append(log_entry)
            logger.warning(log_entry)
        
        # Логируем информацию о системе
        python_version = platform.python_version()
        system_info = f"{platform.system()} {platform.release()} ({platform.processor()})"
        log_entry = f"[{datetime.now().isoformat()}] Системная информация: Python {python_version}, OS: {system_info}"
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Логируем версии библиотек
        try:
            lib_versions = {}
            for lib_name in ["torch", "transformers", "langchain", "ragas", "sentence_transformers", "huggingface_hub", "datasets"]:
                try:
                    lib = importlib.import_module(lib_name)
                    version = getattr(lib, "__version__", "unknown")
                    lib_versions[lib_name] = version
                except (ImportError, AttributeError):
                    lib_versions[lib_name] = "not installed"
            
            log_entry = f"[{datetime.now().isoformat()}] Версии библиотек: {lib_versions}"
            rag_eval_logs[task_id].append(log_entry)
            logger.info(log_entry)
        except Exception as e:
            log_entry = f"[{datetime.now().isoformat()}] Ошибка при получении версий библиотек: {e}"
            rag_eval_logs[task_id].append(log_entry)
            logger.warning(log_entry)
        
        # Проверка запроса на корректность
        if not hasattr(request, 'eval_items'):
            error_message = "Ошибка: Запрос не содержит поля eval_items."
            log_entry = f"[{datetime.now().isoformat()}] {error_message}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            
            rag_eval_progress[task_id]["status"] = "error"
            rag_eval_progress[task_id]["error"] = error_message
            rag_eval_progress[task_id]["completion_time"] = datetime.now().isoformat()
            rag_eval_progress[task_id]["execution_time_sec"] = (datetime.now() - start_time).total_seconds()
            return
        
        # Проверка наличия элементов в датасете
        if not request.eval_items or len(request.eval_items) == 0:
            error_message = "Ошибка: Датасет для оценки пуст. Поле eval_items не содержит элементов."
            log_entry = f"[{datetime.now().isoformat()}] {error_message}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            
            rag_eval_progress[task_id]["status"] = "error"
            rag_eval_progress[task_id]["error"] = error_message
            rag_eval_progress[task_id]["completion_time"] = datetime.now().isoformat()
            rag_eval_progress[task_id]["execution_time_sec"] = (datetime.now() - start_time).total_seconds()
            return
        
        # Логируем конфигурацию
        ollama_model = os.getenv('OLLAMA_MODEL', 'mistral:7b-instruct')
        embedding_model = os.getenv('EMBEDDING_MODEL', 'intfloat/multilingual-e5-base')
        log_entry = f"[{datetime.now().isoformat()}] Модель Ollama: {ollama_model}"
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        log_entry = f"[{datetime.now().isoformat()}] Модель эмбеддингов: {embedding_model}"
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Импортируем здесь, чтобы избежать циклических импортов
        log_entry = f"[{datetime.now().isoformat()}] Импорт необходимых модулей..."
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        from scripts.evaluate_rag import RAGEvaluator
        from app.main import get_rag_service
        
        # Получаем инстанс сервиса RAG
        log_entry = f"[{datetime.now().isoformat()}] Получение RAG-сервиса..."
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        try:
            rag_service = get_rag_service()
            log_entry = f"[{datetime.now().isoformat()}] RAG-сервис успешно получен, тип: {type(rag_service)}"
            rag_eval_logs[task_id].append(log_entry)
            logger.info(log_entry)
        except Exception as e:
            log_entry = f"[{datetime.now().isoformat()}] Ошибка при получении RAG-сервиса: {e}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            rag_eval_progress[task_id] = {
                "status": "error",
                "error": f"Ошибка при инициализации RAG-сервиса: {str(e)}",
                "completion_time": datetime.now().isoformat(),
                "execution_time_sec": (datetime.now() - start_time).total_seconds()
            }
            return
        
        # Создаем новый оценщик с использованием новой реализации
        log_entry = f"[{datetime.now().isoformat()}] Инициализация RAG-оценщика..."
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        try:
            evaluator = RAGEvaluator(
                ollama_model_name=ollama_model,
                embedding_model_name=embedding_model
            )
            log_entry = f"[{datetime.now().isoformat()}] RAG-оценщик успешно инициализирован"
            rag_eval_logs[task_id].append(log_entry)
            logger.info(log_entry)
        except Exception as e:
            log_entry = f"[{datetime.now().isoformat()}] Ошибка при инициализации RAG-оценщика: {e}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            rag_eval_progress[task_id] = {
                "status": "error",
                "error": f"Ошибка при инициализации RAG-оценщика: {str(e)}",
                "completion_time": datetime.now().isoformat(),
                "execution_time_sec": (datetime.now() - start_time).total_seconds()
            }
            return
        
        # Обновляем статус
        rag_eval_progress[task_id]["status"] = "preparing_dataset"
        rag_eval_progress[task_id]["progress"] = 10
        
        log_entry = f"[{datetime.now().isoformat()}] Подготовка датасета для оценки..."
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Собираем данные для оценки
        eval_data = []
        for i, item in enumerate(request.eval_items):
            try:
                eval_item = {
                    "question": item.question,
                    "answer": item.answer if hasattr(item, "answer") else "",
                    "ground_truth": item.ground_truth
                }
                # Проверка корректности данных
                if not eval_item["question"] or not eval_item["ground_truth"]:
                    log_entry = f"[{datetime.now().isoformat()}] Предупреждение: элемент {i} датасета содержит пустые поля: вопрос={bool(eval_item['question'])}, ответ={bool(eval_item['ground_truth'])}"
                    rag_eval_logs[task_id].append(log_entry)
                    logger.warning(log_entry)
                    continue
                eval_data.append(eval_item)
            except Exception as e:
                log_entry = f"[{datetime.now().isoformat()}] Ошибка при обработке элемента {i}: {e}"
                rag_eval_logs[task_id].append(log_entry)
                logger.error(log_entry)
                continue
        
        log_entry = f"[{datetime.now().isoformat()}] Подготовлено {len(eval_data)} элементов для оценки"
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Проверяем, не пустой ли датасет
        if not eval_data:
            error_message = "Ошибка: после обработки и фильтрации датасет не содержит валидных элементов для оценки."
            log_entry = f"[{datetime.now().isoformat()}] {error_message}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            
            rag_eval_progress[task_id]["status"] = "error"
            rag_eval_progress[task_id]["error"] = error_message
            rag_eval_progress[task_id]["completion_time"] = datetime.now().isoformat()
            rag_eval_progress[task_id]["execution_time_sec"] = (datetime.now() - start_time).total_seconds()
            return
        
        # Логируем первый элемент для отладки
        log_entry = f"[{datetime.now().isoformat()}] Пример данных для оценки: {json.dumps(eval_data[0], ensure_ascii=False)}"
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Асинхронно запускаем оценку
        import asyncio
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        # Обновляем статус
        rag_eval_progress[task_id]["status"] = "evaluating"
        rag_eval_progress[task_id]["progress"] = 30
        
        log_entry = f"[{datetime.now().isoformat()}] Запуск процесса оценки RAG..."
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Запуск оценки
        eval_start_time = datetime.now()
        try:
            results = loop.run_until_complete(evaluator.run_evaluation(eval_data))
            loop.close()
            
            # Проверяем наличие ошибки в результатах
            if isinstance(results, dict) and "error" in results:
                error_message = results["error"]
                log_entry = f"[{datetime.now().isoformat()}] Ошибка при выполнении оценки: {error_message}"
                rag_eval_logs[task_id].append(log_entry)
                logger.error(log_entry)
                
                rag_eval_progress[task_id]["status"] = "error"
                rag_eval_progress[task_id]["error"] = error_message
                if "_timing" in results:
                    rag_eval_progress[task_id]["timing"] = results["_timing"]
                rag_eval_progress[task_id]["completion_time"] = datetime.now().isoformat()
                rag_eval_progress[task_id]["execution_time_sec"] = (datetime.now() - eval_start_time).total_seconds()
                return
                
            eval_time = (datetime.now() - eval_start_time).total_seconds()
            log_entry = f"[{datetime.now().isoformat()}] Оценка успешно выполнена за {eval_time:.2f} сек. на {len(eval_data)} элементах датасета"
            rag_eval_logs[task_id].append(log_entry)
            logger.info(log_entry)
        except Exception as e:
            tb = traceback.format_exc()
            log_entry = f"[{datetime.now().isoformat()}] Критическая ошибка при выполнении оценки: {e}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            
            log_entry = f"[{datetime.now().isoformat()}] Трассировка: {tb}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            
            rag_eval_progress[task_id]["status"] = "error" 
            rag_eval_progress[task_id]["error"] = f"Критическая ошибка: {str(e)}"
            rag_eval_progress[task_id]["completion_time"] = datetime.now().isoformat()
            rag_eval_progress[task_id]["execution_time_sec"] = (datetime.now() - eval_start_time).total_seconds()
        
        # Обновляем статус
        rag_eval_progress[task_id]["status"] = "saving_results"
        rag_eval_progress[task_id]["progress"] = 80
        
        log_entry = f"[{datetime.now().isoformat()}] Оценка завершена за {eval_time:.2f} сек. Сохраняю результаты..."
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Логируем метрики
        log_entry = f"[{datetime.now().isoformat()}] Полученные метрики: {json.dumps(results, ensure_ascii=False)}"
        rag_eval_logs[task_id].append(log_entry)
        logger.info(log_entry)
        
        # Сохраняем результаты
        try:
            description = getattr(request, 'description', None) or f"Evaluation {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            output_path = evaluator.save_results(results, eval_data, description)
            
            log_entry = f"[{datetime.now().isoformat()}] Результаты сохранены: {output_path}"
            rag_eval_logs[task_id].append(log_entry)
            logger.info(log_entry)
            
            # Читаем json-файл с результатами и кладём в прогресс
            with open(output_path, "r", encoding="utf-8") as f:
                eval_json = json.load(f)
            
            # Готово, обновляем статус
            total_time = (datetime.now() - start_time).total_seconds()
            
            rag_eval_progress[task_id]["progress"] = 100
            rag_eval_progress[task_id]["status"] = "done"
            rag_eval_progress[task_id]["result"] = eval_json
            rag_eval_progress[task_id]["completion_time"] = datetime.now().isoformat()
            rag_eval_progress[task_id]["execution_time_sec"] = total_time
            
            log_entry = f"[{datetime.now().isoformat()}] Оценка RAG завершена успешно за {total_time:.2f} сек."
            rag_eval_logs[task_id].append(log_entry)
            logger.info(log_entry)
            
            # Выводим краткое резюме результатов
            metrics_summary = []
            for metric_name, value in eval_json.get("metrics", {}).items():
                if isinstance(value, (int, float)):
                    metrics_summary.append(f"{metric_name}: {value:.4f}")
            
            if metrics_summary:
                log_entry = f"[{datetime.now().isoformat()}] Результаты метрик: {', '.join(metrics_summary)}"
                rag_eval_logs[task_id].append(log_entry)
                logger.info(log_entry)
                
            # Добавим вывод информации о каждом примере
            examples_report = eval_json.get("example_reports", [])
            if examples_report:
                log_entry = f"[{datetime.now().isoformat()}] Детальный отчет по каждому примеру:"
                rag_eval_logs[task_id].append(log_entry)
                logger.info(log_entry)
                
                for i, example in enumerate(examples_report):
                    log_entry = f"[{datetime.now().isoformat()}] Пример #{i+1}: "
                    log_entry += f"Вопрос: \"{example.get('question', '')}\" | "
                    log_entry += f"Ответ ({example.get('response_source', 'unknown')}): \"{example.get('response', '')[:100]}...\""
                    rag_eval_logs[task_id].append(log_entry)
                    logger.info(log_entry)
        except Exception as e:
            tb = traceback.format_exc()
            log_entry = f"[{datetime.now().isoformat()}] Ошибка при сохранении результатов: {e}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            
            log_entry = f"[{datetime.now().isoformat()}] Трассировка: {tb}"
            rag_eval_logs[task_id].append(log_entry)
            logger.error(log_entry)
            
            rag_eval_progress[task_id]["status"] = "error" 
            rag_eval_progress[task_id]["error"] = f"Ошибка сохранения результатов: {str(e)}"
            rag_eval_progress[task_id]["raw_results"] = results
            rag_eval_progress[task_id]["completion_time"] = datetime.now().isoformat()
            rag_eval_progress[task_id]["execution_time_sec"] = (datetime.now() - start_time).total_seconds()
        
    except Exception as e:
        tb = traceback.format_exc()
        error_time = (datetime.now() - start_time).total_seconds()
        
        error_message = f"Ошибка оценки RAG: {str(e)}"
        log_entry = f"[{datetime.now().isoformat()}] {error_message}"
        rag_eval_logs[task_id].append(log_entry)
        logger.error(log_entry)
        
        log_entry = f"[{datetime.now().isoformat()}] Трассировка:\n{tb}"
        rag_eval_logs[task_id].append(log_entry)
        logger.error(log_entry)
        
        rag_eval_progress[task_id] = {
            "status": "error", 
            "error": str(e), 
            "traceback": tb,
            "start_time": start_time.isoformat(),
            "completion_time": datetime.now().isoformat(),
            "execution_time_sec": error_time
        }

@app.post("/api/evaluate-rag/async")
async def evaluate_rag_async(request: RAGEvaluationRequest, background_tasks: BackgroundTasks, current_user: UserModel = Depends(get_current_user), db: Session = Depends(get_db)):
    task_id = str(uuid.uuid4())
    background_tasks.add_task(run_rag_evaluation_task, task_id, request, db, current_user.id)
    return {"task_id": task_id}

@app.get("/api/evaluate-rag/status/{task_id}")
async def get_rag_eval_status(task_id: str):
    status = rag_eval_progress.get(task_id)
    if not status:
        raise HTTPException(status_code=404, detail="Task not found")
    
    # Обработка специальных значений float перед сериализацией в JSON
    def clean_float_values(obj):
        if isinstance(obj, dict):
            return {k: clean_float_values(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [clean_float_values(item) for item in obj]
        elif isinstance(obj, float):
            # Преобразование inf, -inf, nan в строки или обычные числа
            if math.isinf(obj):
                return 1.0 if obj > 0 else -1.0  # Заменяем inf на 1.0, -inf на -1.0
            elif math.isnan(obj):
                return 0.0  # Заменяем NaN на 0.0
            return obj
        else:
            return obj
    
    # Очистка специальных значений в статусе
    cleaned_status = clean_float_values(status)
    return cleaned_status

@app.get("/api/evaluate-rag/logs/{task_id}")
async def get_rag_eval_logs(task_id: str):
    return {"logs": rag_eval_logs.get(task_id, [])}

@app.delete("/documents/{document_id}")
async def delete_document(
    document_id: int,
    current_user: UserModel = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Удаляет документ пользователя по id (и из БД, и из индекса).
    """
    doc = db.query(DocumentModel).filter(
        DocumentModel.id == document_id,
        DocumentModel.user_id == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не найден")
    # Удаляем из индекса (если нужно — можно реализовать удаление чанков по doc_id)
    rag = get_rag_service()
    # В Qdrant нет удаления по doc_id, но можно реализовать через фильтр (TODO: если потребуется)
    # Сейчас просто очищаем весь индекс, если нужно — доработать под удаление отдельных чанков
    # rag.index.delete_by_doc_id(str(doc.id))
    db.delete(doc)
    db.commit()
    return {"status": "success", "message": f"Документ {document_id} удалён"}

@app.post("/index/clear")
async def clear_index(
    current_user: UserModel = Depends(get_current_user)
):
    """
    Полностью очищает векторный индекс (Qdrant).
    """
    rag = get_rag_service()
    rag.index.clear()
    return {"status": "success", "message": "Индекс очищен"}

@app.get("/documents/{document_id}")
async def get_document_detail(document_id: int, current_user: UserModel = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(DocumentModel).filter(DocumentModel.id == document_id, DocumentModel.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не найден")
    return {
        "id": doc.id,
        "uuid": doc.uuid,
        "title": doc.title,
        "source": doc.source,
        "created_at": doc.created_at,
        "updated_at": doc.updated_at,
        "status": doc.status,
        "file_name": doc.file_name,
        "file_size": doc.file_size,
        "chunks_count": doc.chunks_count,
        "error_message": doc.error_message,
        "content_length": len(doc.content) if doc.content else 0
    }

@app.get("/api/evaluation/last_results")
async def get_last_rag_evaluation_results():
    import os
    import json
    from pathlib import Path
    result_dir = Path("app/evaluation/results")
    files = sorted(result_dir.glob("rag_eval_results_*.json"), key=os.path.getmtime, reverse=True)
    if not files:
        return JSONResponse(status_code=404, content={"detail": "Нет сохранённых результатов оценки"})
    with open(files[0], "r", encoding="utf-8") as f:
        data = json.load(f)
    return data

@app.post("/test/init-db")
def test_init_db():
    create_tables()
    return {"status": "ok"}

@app.get("/api/evaluation/history")
async def get_rag_evaluation_history():
    """
    Возвращает список всех запусков оценки RAG (метрики, датасет, timestamp, имя файла, summary, trace_id если есть).
    """
    import os
    import json
    from pathlib import Path
    result_dir = Path("app/evaluation/results")
    files = sorted(result_dir.glob("rag_eval_results_*.json"), key=os.path.getmtime, reverse=True)
    history = []
    for f in files:
        try:
            with open(f, "r", encoding="utf-8") as fp:
                data = json.load(fp)
                history.append({
                    "metrics": data.get("metrics"),
                    "dataset_name": data.get("dataset_name"),
                    "timestamp": data.get("timestamp"),
                    "filename": f.name,
                    "summary": data.get("summary"),
                    "trace_id": data.get("meta", {}).get("trace_id")
                })
        except Exception as e:
            continue
    return history

@app.get("/api/evaluation/download/{filename}")
async def download_rag_evaluation_report(filename: str):
    """
    Скачивание json-отчёта по имени файла (только из results, без выхода за пределы папки).
    """
    from pathlib import Path
    import os
    safe_dir = Path("app/evaluation/results").resolve()
    file_path = (safe_dir / filename).resolve()
    if not str(file_path).startswith(str(safe_dir)) or not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")
    return FileResponse(str(file_path), media_type="application/json", filename=filename)

from fastapi import UploadFile, File, Form
import shutil

@app.post("/api/evaluation/upload_dataset")
async def upload_evaluation_dataset(
    dataset_name: str = Form(...),
    file: UploadFile = File(None),
    url: str = Form(None)
):
    """
    Загрузка датасета для оценки (json-файл или по URL). Имя файла = {dataset_name}.json
    """
    from pathlib import Path
    import requests
    import json
    datasets_dir = Path("app/evaluation/datasets")
    datasets_dir.mkdir(parents=True, exist_ok=True)
    target_path = datasets_dir / f"{dataset_name}.json"
    if file:
        with open(target_path, "wb") as out_f:
            shutil.copyfileobj(file.file, out_f)
    elif url:
        r = requests.get(url, timeout=20)
        r.raise_for_status()
        with open(target_path, "wb") as out_f:
            out_f.write(r.content)
    else:
        raise HTTPException(status_code=400, detail="Нужно передать файл или url")
    # Валидация структуры (ожидаем список dict с question/answer/ground_truth)
    try:
        with open(target_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        assert isinstance(data, list)
        for item in data:
            assert isinstance(item, dict)
            assert "question" in item and "answer" in item and "ground_truth" in item
    except Exception as e:
        target_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"Некорректный формат датасета: {e}")
    return {"status": "ok", "dataset_name": dataset_name, "count": len(data)}

# --- RAGAS Evaluation Endpoints ---
ragas_eval_progress: Dict[str, Dict] = {}
ragas_eval_logs: Dict[str, list] = {} # Отдельные логи для RAGAS

def run_ragas_evaluation_task(
    task_id: str, 
    request_data: RagasEvaluationRequest, 
    db: Session, 
    user_id: int 
):
    import traceback
    import json
    global ragas_eval_progress, ragas_eval_logs
    
    ragas_eval_progress[task_id] = {"status": "starting", "progress": 0, "description": request_data.description or "RAGAS Evaluation"}
    ragas_eval_logs[task_id] = [f"[{datetime.utcnow().isoformat()}] Starting RAGAS evaluation task ID: {task_id}"]
    
    try:
        logger.info(f"RAGAS Evaluation Task {task_id}: Initializing services.")
        ragas_eval_logs[task_id].append(f"[{datetime.utcnow().isoformat()}] Initializing RAG and Ollama services.")
        
        # Используем новый RAGEvaluator из evaluate_rag.py
        from scripts.evaluate_rag import RAGEvaluator
        
        # Создаем оценщик напрямую, без использования RAGService
        evaluator = RAGEvaluator(
            ollama_model_name=os.getenv("OLLAMA_MODEL", "mistral:7b-instruct"),
            embedding_model_name=os.getenv("EMBEDDING_MODEL", "intfloat/multilingual-e5-base")
        )
        
        ragas_eval_progress[task_id]["status"] = "preparing_data"
        ragas_eval_progress[task_id]["progress"] = 10
        ragas_eval_logs[task_id].append(f"[{datetime.utcnow().isoformat()}] Preparing data for RAGAS evaluation...")
        
        # Преобразуем данные из запроса в формат для оценщика
        eval_data = []
        for item in request_data.eval_items:
            eval_data.append({
                "question": item.question,
                "answer": item.answer if hasattr(item, "answer") else "",
                "ground_truth": item.ground_truth
            })
        
        # Создаем новый event loop для асинхронного вызова
        import asyncio
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        # Запускаем оценку
        ragas_eval_logs[task_id].append(f"[{datetime.utcnow().isoformat()}] Running RAGAS evaluation...")
        ragas_eval_progress[task_id]["progress"] = 30
        
        results = loop.run_until_complete(evaluator.run_evaluation(eval_data))
        loop.close()

        ragas_eval_progress[task_id]["status"] = "processing_results"
        ragas_eval_progress[task_id]["progress"] = 90
        ragas_eval_logs[task_id].append(f"[{datetime.utcnow().isoformat()}] RAGAS evaluation completed. Processing results.")

        if "error" in results:
            logger.error(f"RAGAS Evaluation Task {task_id}: Error in RAGAS results - {results['error']}")
            ragas_eval_logs[task_id].append(f"[{datetime.utcnow().isoformat()}] ERROR in RAGAS results: {results['error']}")
            ragas_eval_progress[task_id] = {"status": "error", "error": results['error']}
        else:
            # Сохраняем результаты в файл
            output_path = evaluator.save_results(results, eval_data, "custom")
            
            ragas_eval_logs[task_id].append(f"[{datetime.utcnow().isoformat()}] Results saved to {output_path}")
            ragas_eval_progress[task_id]["status"] = "done"
            ragas_eval_progress[task_id]["progress"] = 100
            ragas_eval_progress[task_id]["result_file"] = str(output_path)
            
            # Формируем summary из метрик
            metrics_avg = {}
            for metric_name, value in results.items():
                if isinstance(value, (int, float)):
                    metrics_avg[metric_name] = value
            
            ragas_eval_progress[task_id]["results_summary"] = metrics_avg

        ragas_eval_logs[task_id].append(f"[{datetime.utcnow().isoformat()}] RAGAS evaluation task finished.")

    except Exception as e:
        tb = traceback.format_exc()
        logger.error(f"""RAGAS Evaluation Task {task_id} failed: {e}
Traceback:
{tb}""")
        ragas_eval_logs[task_id].append(f"""[{datetime.utcnow().isoformat()}] CRITICAL ERROR: {e}
Traceback:
{tb}""")
        ragas_eval_progress[task_id] = {"status": "error", "error": str(e), "traceback": tb}

@app.post("/api/evaluate-ragas/async", summary="Запустить оценку RAG с использованием RAGAS (асинхронно)")
async def evaluate_ragas_async(
    request: RagasEvaluationRequest, 
    background_tasks: BackgroundTasks, 
    current_user: UserModel = Depends(get_current_user), # Добавляем current_user
    db: Session = Depends(get_db) # Добавляем db
):
    if not request.eval_items:
        raise HTTPException(status_code=400, detail="eval_items не может быть пустым")

    task_id = str(uuid.uuid4())
    # Передаем request.dict() или конкретные поля, так как request - Pydantic модель
    # background_tasks.add_task ожидает аргументы, которые можно передать в функцию
    background_tasks.add_task(
        run_ragas_evaluation_task, 
        task_id, 
        request, # Передаем всю Pydantic модель
        db, # Передаем сессию БД
        current_user.id # Передаем ID пользователя
    )
    return {"task_id": task_id, "message": "RAGAS evaluation task started."}

@app.get("/api/evaluate-ragas/status/{task_id}", summary="Получить статус задачи оценки RAGAS")
async def get_ragas_eval_status(task_id: str):
    status = ragas_eval_progress.get(task_id)
    if not status:
        raise HTTPException(status_code=404, detail="Task not found")
    return status

@app.get("/api/evaluate-ragas/logs/{task_id}", summary="Получить логи задачи оценки RAGAS")
async def get_ragas_eval_logs(task_id: str):
    logs = ragas_eval_logs.get(task_id)
    if logs is None: # Проверяем на None, так как пустой список - валидное значение
        raise HTTPException(status_code=404, detail="Task not found or no logs available")
    return {"logs": logs}

@app.get("/api/evaluate-ragas/results", summary="Получить список всех сохраненных результатов RAGAS")
async def get_all_ragas_evaluation_results(skip: int = 0, limit: int = 10):
    results_dir = Path("app/evaluation/results/ragas")
    if not results_dir.exists():
        return []
    
    result_files = sorted(
        [f for f in results_dir.glob("ragas_eval_results_*.json") if f.is_file()],
        key=os.path.getmtime,
        reverse=True
    )
    
    history = []
    for f_path in result_files[skip : skip + limit]:
        try:
            with open(f_path, "r", encoding="utf-8") as f:
                data = json.load(f) # data is a dict of lists like {'metric': [s1,s2,...]}
                # Извлекаем ключевые метрики для summary путем усреднения
                metrics_summary = {}
                if isinstance(data, dict):
                    for metric_name, scores_list in data.items():
                        if metric_name not in ['question', 'answer', 'contexts', 'ground_truth', 'episode_done'] and isinstance(scores_list, list):
                            numeric_scores = [s for s in scores_list if isinstance(s, (int, float))]
                            if numeric_scores:
                                metrics_summary[metric_name] = sum(numeric_scores) / len(numeric_scores)
                            else:
                                metrics_summary[metric_name] = None
                                
                history.append({
                    "filename": f_path.name,
                    "timestamp": datetime.fromtimestamp(f_path.stat().st_mtime).isoformat(),
                    "metrics_summary": metrics_summary 
                    # Можно добавить больше информации, если она есть в файле
                })
        except Exception as e:
            logger.warning(f"Could not read or parse RAGAS result file {f_path.name}: {e}")
            history.append({
                "filename": f_path.name,
                "error": "Failed to load or parse."
            })
            
    return history

@app.get("/api/evaluate-ragas/results/{filename}", summary="Скачать конкретный файл результатов RAGAS")
async def download_ragas_evaluation_report(filename: str):
    safe_dir = Path("app/evaluation/results/ragas").resolve()
    file_path = (safe_dir / filename).resolve()
    
    if not str(file_path).startswith(str(safe_dir)) or not file_path.is_file():
        raise HTTPException(status_code=404, detail="File not found or access denied.")
    
    return FileResponse(str(file_path), media_type="application/json", filename=filename)

# --- End RAGAS Evaluation Endpoints ---

# Эндпоинт для получения содержимого датасета
@app.get("/api/evaluation/datasets/{dataset_name}")
async def get_evaluation_dataset(dataset_name: str, limit: int = 20):
    """
    Возвращает содержимое датасета для оценки по имени
    """
    try:
        # Получаем список доступных датасетов
        datasets = get_builtin_datasets()
        
        if dataset_name not in datasets:
            raise HTTPException(status_code=404, detail=f"Датасет {dataset_name} не найден")
        
        dataset_path = datasets[dataset_name]
        if not os.path.exists(dataset_path):
            raise HTTPException(status_code=404, detail=f"Файл датасета {dataset_name} не найден")
        
        # Загружаем датасет
        with open(dataset_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Проверяем формат
        if isinstance(data, dict) and 'eval_items' in data:
            items = data['eval_items']
        elif isinstance(data, list):
            items = data
        else:
            raise HTTPException(status_code=400, detail=f"Неподдерживаемый формат данных в {dataset_name}")
        
        # Возвращаем метаданные о датасете и ограниченное количество элементов
        return {
            "name": dataset_name,
            "total_items": len(items),
            "path": str(dataset_path),
            "format": "list" if isinstance(data, list) else "dict",
            "example_fields": list(items[0].keys()) if items else [],
            "items": items[:limit]  # Ограничиваем количество элементов
        }
    except Exception as e:
        logger.error(f"Ошибка при получении датасета {dataset_name}: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка при получении датасета: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 